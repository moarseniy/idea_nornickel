from __future__ import annotations

import csv
import io
import re
import warnings
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


SUPPORTED_EXTENSIONS = {".txt", ".md", ".csv", ".pdf", ".docx", ".xlsx", ".xlsm", ".png", ".jpg", ".jpeg"}


@dataclass
class ParsedDocument:
    text: str
    metadata: dict[str, object] = field(default_factory=dict)
    vision_images: list["VisionImage"] = field(default_factory=list)


@dataclass
class VisionImage:
    label: str
    data: bytes
    content_type: str = "image/png"
    reason: str = "image"
    page_no: int | None = None


def safe_filename(filename: str) -> str:
    stem = Path(filename).name.strip() or "document"
    return re.sub(r"[^A-Za-zА-Яа-я0-9._ -]+", "_", stem)[:180]


def parse_document(
    filename: str,
    content_type: str | None,
    data: bytes,
    max_chars: int,
    *,
    pdf_ocr_enabled: bool = True,
    pdf_ocr_languages: list[str] | None = None,
    pdf_ocr_model_dir: Path | None = None,
    pdf_ocr_min_chars: int = 12,
    pdf_text_layer_min_chars: int = 16,
    pdf_render_dpi: int = 160,
    pdf_vision_max_pages: int = 8,
) -> ParsedDocument:
    extension = Path(filename).suffix.lower()
    metadata: dict[str, object] = {
        "filename": filename,
        "content_type": content_type or "application/octet-stream",
        "extension": extension,
        "bytes": len(data),
    }
    selected_ocr_languages = normalize_ocr_languages(pdf_ocr_languages, default=["ru", "en", "ch_sim"])
    metadata["requested_ocr_languages"] = selected_ocr_languages
    vision_images: list[VisionImage] = []

    if extension in {".txt", ".md"}:
        text = _decode_text(data)
    elif extension == ".csv":
        text = _parse_csv(data)
    elif extension == ".pdf":
        text, pdf_meta, vision_images = _parse_pdf(
            filename=filename,
            data=data,
            ocr_enabled=pdf_ocr_enabled,
            ocr_languages=selected_ocr_languages,
            ocr_model_dir=pdf_ocr_model_dir,
            ocr_min_chars=pdf_ocr_min_chars,
            text_layer_min_chars=pdf_text_layer_min_chars,
            render_dpi=pdf_render_dpi,
            vision_max_pages=pdf_vision_max_pages,
        )
        metadata.update(pdf_meta)
    elif extension == ".docx":
        text, doc_meta = _parse_docx(data)
        metadata.update(doc_meta)
    elif extension in {".xlsx", ".xlsm"}:
        text, xlsx_meta = _parse_xlsx(data)
        metadata.update(xlsx_meta)
    elif extension in {".png", ".jpg", ".jpeg"}:
        text = f"Изображение или схема: {filename}. Текстовое содержание будет дополнено vision-анализом при наличии OpenAI API."
        metadata["image_only"] = True
        vision_images.append(
            VisionImage(
                label=filename,
                data=data,
                content_type=content_type if content_type and content_type.startswith("image/") else _image_content_type(filename),
                reason="standalone_image",
            )
        )
    else:
        text = _decode_text(data)
        metadata["warning"] = "Формат не распознан, выполнена попытка текстового декодирования."

    text = _normalize_text(text)
    if len(text) > max_chars:
        metadata["truncated"] = True
        metadata["original_chars"] = len(text)
        text = text[:max_chars]
    metadata["chars"] = len(text)
    metadata.update(_language_metadata(text))
    return ParsedDocument(text=text, metadata=metadata, vision_images=vision_images)


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1251", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _parse_csv(data: bytes) -> str:
    decoded = _decode_text(data)
    sample = decoded[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample)
    except csv.Error:
        dialect = csv.excel
    rows = []
    for idx, row in enumerate(csv.reader(io.StringIO(decoded), dialect=dialect)):
        if idx >= 500:
            rows.append("... CSV обрезан после 500 строк ...")
            break
        rows.append(" | ".join(cell.strip() for cell in row if cell is not None))
    return "\n".join(rows)


def _parse_pdf(
    filename: str,
    data: bytes,
    ocr_enabled: bool,
    ocr_languages: list[str],
    ocr_model_dir: Path | None,
    ocr_min_chars: int,
    text_layer_min_chars: int,
    render_dpi: int,
    vision_max_pages: int,
) -> tuple[str, dict[str, object], list[VisionImage]]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    fitz_document: Any | None = None
    chunks: list[str] = []
    vision_images: list[VisionImage] = []
    page_modes: list[dict[str, object]] = []
    counters = {"text_layer": 0, "ocr": 0, "vision": 0, "blank": 0, "failed": 0}
    parsed_pages = min(len(reader.pages), 80)

    for page_index, page in enumerate(reader.pages[:parsed_pages]):
        page_no = page_index + 1
        page_text = _normalize_text(page.extract_text() or "")
        has_images = _pdf_page_has_images(page)
        if page_text and (len(page_text) >= max(1, text_layer_min_chars) or not has_images):
            chunks.append(f"[стр. {page_no} | text-layer]\n{page_text}")
            counters["text_layer"] += 1
            page_modes.append({"page": page_no, "mode": "text_layer", "chars": len(page_text)})
            continue

        rendered_page, render_error = _render_pdf_page(data, page_index, render_dpi, fitz_document)
        if rendered_page is None:
            if page_text:
                chunks.append(f"[стр. {page_no} | short text-layer]\n{page_text}")
                counters["text_layer"] += 1
                page_modes.append(
                    {
                        "page": page_no,
                        "mode": "short_text_layer",
                        "chars": len(page_text),
                        "render_error": render_error or "unknown",
                    }
                )
            else:
                counters["failed"] += 1
                page_modes.append({"page": page_no, "mode": "render_failed", "reason": render_error or "unknown"})
            continue
        fitz_document = rendered_page[1]
        image_bytes = rendered_page[0]

        ocr_text = ""
        ocr_error = ""
        if ocr_enabled:
            ocr_text, ocr_error = _ocr_image(image_bytes, ocr_languages, ocr_model_dir, min_chars=ocr_min_chars)
            ocr_text = _normalize_text(ocr_text)

        if len(ocr_text) >= max(1, ocr_min_chars):
            chunks.append(f"[стр. {page_no} | OCR]\n{ocr_text}")
            counters["ocr"] += 1
            page_modes.append({"page": page_no, "mode": "ocr", "chars": len(ocr_text)})
            continue

        if not _image_has_visual_content(image_bytes):
            counters["blank"] += 1
            page_modes.append({"page": page_no, "mode": "blank_render", "ocr_chars": len(ocr_text)})
            continue

        if len(vision_images) < max(0, vision_max_pages):
            vision_images.append(
                VisionImage(
                    label=f"{filename} | стр. {page_no}",
                    data=image_bytes,
                    reason="pdf_image_page_without_ocr_text",
                    page_no=page_no,
                )
            )
            counters["vision"] += 1
            mode: dict[str, object] = {
                "page": page_no,
                "mode": "vision_queued",
                "ocr_chars": len(ocr_text),
                "text_layer_chars": len(page_text),
                "has_pdf_images": has_images,
            }
            if ocr_error:
                mode["ocr_error"] = ocr_error
            page_modes.append(mode)
        else:
            counters["blank"] += 1
            page_modes.append({"page": page_no, "mode": "vision_skipped_limit", "ocr_chars": len(ocr_text)})

    _close_fitz_document(fitz_document)
    meta = {
        "pages": len(reader.pages),
        "parsed_pages": parsed_pages,
        "pdf_parse": counters,
        "pdf_page_modes": page_modes,
        "ocr_enabled": ocr_enabled,
        "ocr_languages": ocr_languages,
        "ocr_language_groups": _ocr_language_groups(ocr_languages),
        "text_layer_min_chars": text_layer_min_chars,
        "vision_queued_pages": len(vision_images),
    }
    return "\n\n".join(chunks), meta, vision_images


def _parse_docx(data: bytes) -> tuple[str, dict[str, object]]:
    from docx import Document

    document = Document(io.BytesIO(data))
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if value:
            chunks.append(value)
    for table_no, table in enumerate(document.tables, start=1):
        chunks.append(f"[таблица {table_no}]")
        for row in table.rows[:250]:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))
    return "\n".join(chunks), {"paragraphs": len(document.paragraphs), "tables": len(document.tables)}


def _parse_xlsx(data: bytes) -> tuple[str, dict[str, object]]:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    chunks: list[str] = []
    sheet_names = workbook.sheetnames
    for sheet_name in sheet_names[:12]:
        sheet = workbook[sheet_name]
        chunks.append(f"[лист: {sheet_name}]")
        for row_no, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            if row_no > 450:
                chunks.append("... лист обрезан после 450 строк ...")
                break
            values = [_cell_to_text(value) for value in row]
            if any(values):
                chunks.append(" | ".join(values))
    workbook.close()
    return "\n".join(chunks), {"sheets": sheet_names, "parsed_sheets": min(len(sheet_names), 12)}


def _pdf_page_has_images(page: Any) -> bool:
    try:
        if getattr(page, "images", None):
            return True
    except Exception:  # noqa: BLE001
        pass

    try:
        resources = _resolve_pdf_object(page.get("/Resources") or {})
        xobjects = _resolve_pdf_object(resources.get("/XObject") or {})
        if not hasattr(xobjects, "values"):
            return False
        for raw_xobject in xobjects.values():
            xobject = _resolve_pdf_object(raw_xobject)
            if xobject.get("/Subtype") == "/Image":
                return True
    except Exception:  # noqa: BLE001
        return False
    return False


def _resolve_pdf_object(value: Any) -> Any:
    if hasattr(value, "get_object"):
        return value.get_object()
    return value


def _render_pdf_page(
    data: bytes,
    page_index: int,
    dpi: int,
    document: Any | None,
) -> tuple[tuple[bytes, Any] | None, str]:
    try:
        import fitz

        if document is None:
            document = fitz.open(stream=data, filetype="pdf")
        page = document.load_page(page_index)
        scale = max(72, dpi) / 72
        pixmap = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
        return (pixmap.tobytes("png"), document), ""
    except Exception as exc:  # noqa: BLE001
        return None, str(exc)


def _close_fitz_document(document: Any | None) -> None:
    if document is None:
        return
    try:
        document.close()
    except Exception:  # noqa: BLE001
        pass


_EASYOCR_READER_CACHE: dict[tuple[tuple[str, ...], str], Any] = {}


def _ocr_image(image_bytes: bytes, languages: list[str], model_dir: Path | None, min_chars: int = 1) -> tuple[str, str]:
    groups = _ocr_language_groups(languages)
    if not groups:
        return "", "No OCR languages configured"

    errors: list[str] = []
    candidates: list[tuple[float, str, list[str]]] = []
    for group in groups:
        text, error = _ocr_image_with_group(image_bytes, group, model_dir)
        normalized_text = _normalize_text(text)
        if normalized_text:
            candidates.append((_ocr_quality_score(normalized_text, group, min_chars), normalized_text, group))
        if error:
            errors.append(f"{'/'.join(group)}: {error}")

    if candidates:
        candidates.sort(key=lambda item: item[0], reverse=True)
        return candidates[0][1], "; ".join(errors)
    return "", "; ".join(errors)


def _ocr_image_with_group(image_bytes: bytes, languages: list[str], model_dir: Path | None) -> tuple[str, str]:
    try:
        import numpy as np
        from PIL import Image

        reader = _get_easyocr_reader(languages, model_dir)
        image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
        with warnings.catch_warnings():
            warnings.filterwarnings("ignore", message="'pin_memory' argument is set as true.*", category=UserWarning)
            result = reader.readtext(np.array(image), detail=0, paragraph=True)
        if isinstance(result, list):
            return "\n".join(str(item) for item in result if str(item).strip()), ""
        return str(result), ""
    except Exception as exc:  # noqa: BLE001
        return "", str(exc)


def _ocr_language_groups(languages: list[str]) -> list[list[str]]:
    normalized = normalize_ocr_languages(languages, default=["ru", "en", "ch_sim"])
    if "en" not in normalized:
        normalized.append("en")

    groups: list[list[str]] = []
    non_chinese = [code for code in normalized if code not in {"ch_sim", "ch_tra"}]
    if non_chinese:
        groups.append(_dedupe_languages(non_chinese))
    for chinese_code in ("ch_sim", "ch_tra"):
        if chinese_code in normalized:
            groups.append(_dedupe_languages([chinese_code, "en"]))
    return groups


def normalize_ocr_languages(languages: list[str] | str | None, default: list[str] | None = None) -> list[str]:
    if isinstance(languages, str):
        raw_languages = re.split(r"[,;\s]+", languages)
    else:
        raw_languages = languages or []
    normalized: list[str] = []
    for language in raw_languages:
        code = _normalize_ocr_language_code(str(language))
        if code and code not in normalized:
            normalized.append(code)
    return normalized or list(default or [])


def _normalize_ocr_language_code(language: str) -> str:
    code = language.strip().lower().replace("-", "_")
    aliases = {
        "eng": "en",
        "english": "en",
        "rus": "ru",
        "russian": "ru",
        "zh": "ch_sim",
        "zh_cn": "ch_sim",
        "cn": "ch_sim",
        "chinese": "ch_sim",
        "chinese_simplified": "ch_sim",
        "zh_tw": "ch_tra",
        "zh_hk": "ch_tra",
        "traditional_chinese": "ch_tra",
    }
    return aliases.get(code, code)


def _dedupe_languages(languages: list[str]) -> list[str]:
    result: list[str] = []
    for language in languages:
        if language not in result:
            result.append(language)
    return result


def _ocr_quality_score(text: str, languages: list[str], min_chars: int) -> float:
    counts = _script_counts(text)
    score = float(len(text))
    if "ru" in languages:
        score += counts["cyrillic"] * 2.5
    if "en" in languages:
        score += counts["latin"] * 1.3
    if any(language in {"ch_sim", "ch_tra"} for language in languages):
        score += counts["han"] * 3.0
    if len(text) < max(1, min_chars):
        score *= 0.25
    unsupported = max(0, counts["letters"] - counts["cyrillic"] - counts["latin"] - counts["han"])
    return score - unsupported


def _image_has_visual_content(image_bytes: bytes) -> bool:
    try:
        from PIL import Image, ImageStat

        image = Image.open(io.BytesIO(image_bytes)).convert("L")
        image.thumbnail((96, 96))
        stat = ImageStat.Stat(image)
        mean = stat.mean[0]
        stddev = stat.stddev[0]
        return mean < 248 or stddev > 3
    except Exception:  # noqa: BLE001
        return True


def _get_easyocr_reader(languages: list[str], model_dir: Path | None) -> Any:
    import easyocr

    normalized_languages = tuple(languages or ["ru", "en"])
    model_path = str(model_dir or "")
    key = (normalized_languages, model_path)
    reader = _EASYOCR_READER_CACHE.get(key)
    if reader is not None:
        return reader
    if model_dir is not None:
        model_dir.mkdir(parents=True, exist_ok=True)
    kwargs: dict[str, object] = {"gpu": False, "verbose": False}
    if model_path:
        kwargs["model_storage_directory"] = model_path
    reader = easyocr.Reader(list(normalized_languages), **kwargs)
    _EASYOCR_READER_CACHE[key] = reader
    return reader


def _language_metadata(text: str) -> dict[str, object]:
    counts = _script_counts(text)
    detected: list[str] = []
    if counts["cyrillic"] >= 3:
        detected.append("ru/cyrillic")
    if counts["latin"] >= 3:
        detected.append("en/latin")
    if counts["han"] >= 2:
        detected.append("zh/han")
    for script, label in (
        ("arabic", "arabic"),
        ("devanagari", "devanagari"),
        ("hiragana_katakana", "ja/kana"),
        ("hangul", "ko/hangul"),
    ):
        if counts[script] >= 2:
            detected.append(label)
    return {
        "detected_languages": detected,
        "script_counts": counts,
        "multilingual": len(detected) > 1,
    }


def _script_counts(text: str) -> dict[str, int]:
    counts = {
        "letters": 0,
        "latin": 0,
        "cyrillic": 0,
        "han": 0,
        "arabic": 0,
        "devanagari": 0,
        "hiragana_katakana": 0,
        "hangul": 0,
    }
    for char in text:
        code = ord(char)
        if char.isalpha():
            counts["letters"] += 1
        if 0x0041 <= code <= 0x007A or 0x00C0 <= code <= 0x024F:
            counts["latin"] += 1
        elif 0x0400 <= code <= 0x052F:
            counts["cyrillic"] += 1
        elif 0x3400 <= code <= 0x9FFF or 0xF900 <= code <= 0xFAFF:
            counts["han"] += 1
        elif 0x0600 <= code <= 0x06FF:
            counts["arabic"] += 1
        elif 0x0900 <= code <= 0x097F:
            counts["devanagari"] += 1
        elif 0x3040 <= code <= 0x30FF:
            counts["hiragana_katakana"] += 1
        elif 0xAC00 <= code <= 0xD7AF:
            counts["hangul"] += 1
    return counts


def _cell_to_text(value: object) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    return re.sub(r"\s+", " ", text)


def _normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _image_content_type(filename: str) -> str:
    extension = Path(filename).suffix.lower()
    if extension in {".jpg", ".jpeg"}:
        return "image/jpeg"
    return "image/png"
